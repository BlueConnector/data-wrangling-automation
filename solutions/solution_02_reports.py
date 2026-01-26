"""
Exercise 2 Solution: Automated Report Generation
Complete implementation of an automated reporting workflow

This solution demonstrates:
1. Data loading and processing with pandas
2. Metric calculation (MoM, YoY growth)
3. Rule-based insight generation
4. Excel report creation with openpyxl
5. Word document generation with python-docx
6. Bonus: Interactive dashboard with Plotly
"""

import pandas as pd
import openpyxl
from openpyxl.chart import LineChart, Reference
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.shared import Pt, RGBColor
from datetime import datetime
import os


def load_monthly_data():
    """
    Load the monthly statistics data from CSV
    
    Returns:
        DataFrame: The loaded data
    """
    try:
        data = pd.read_csv('data/sample_monthly_data.csv')
        print(f"✓ Loaded {len(data)} indicators from CSV")
        print(f"  Indicators: {', '.join(data['indicator_code'].tolist())}")
        return data
    except FileNotFoundError:
        print("✗ Error: data/sample_monthly_data.csv not found")
        print("Make sure you're running from the repository root directory")
        return None


def calculate_metrics(data):
    """
    Calculate key metrics from the data
    
    Args:
        data (DataFrame): The monthly statistics data
        
    Returns:
        dict: Dictionary of calculated metrics
    """
    metrics = {}
    
    print("\nCalculating metrics...")
    
    for idx, row in data.iterrows():
        indicator_code = row['indicator_code']
        
        # Extract values
        current_value = row['value']
        previous_month = row['previous_month_value']
        previous_year = row['previous_year_value']
        
        # Calculate month-over-month growth
        # Formula: ((current - previous) / previous) * 100
        if previous_month and previous_month != 0:
            mom_growth = ((current_value - previous_month) / previous_month) * 100
        else:
            mom_growth = None
        
        # Calculate year-over-year growth
        # Formula: ((current - previous_year) / previous_year) * 100
        if previous_year and previous_year != 0:
            yoy_growth = ((current_value - previous_year) / previous_year) * 100
        else:
            yoy_growth = None
        
        # Determine trend direction
        # Threshold: 1% for "stable"
        if mom_growth is not None:
            if mom_growth > 1.0:
                trend = "increasing"
            elif mom_growth < -1.0:
                trend = "decreasing"
            else:
                trend = "stable"
        else:
            trend = "unknown"
        
        metrics[indicator_code] = {
            'name': row['indicator_name'],
            'code': indicator_code,
            'current_value': current_value,
            'previous_month_value': previous_month,
            'previous_year_value': previous_year,
            'unit': row['unit'],
            'period': row['period'],
            'mom_growth': mom_growth,
            'yoy_growth': yoy_growth,
            'trend': trend
        }
        
        print(f"  {indicator_code}: MoM={mom_growth:.2f}% YoY={yoy_growth:.2f}% Trend={trend}")
    
    return metrics


def generate_insights(metrics):
    """
    Generate insights using rule-based conditional logic
    
    Args:
        metrics (dict): Dictionary of calculated metrics
        
    Returns:
        list: List of insight strings
    """
    insights = []
    
    print("\nGenerating insights...")
    
    for code, data in metrics.items():
        mom = data['mom_growth']
        yoy = data['yoy_growth']
        name = data['name']
        
        # Rule 1: Significant month-over-month growth (>5%)
        if mom is not None and mom > 5:
            insight = f"{name} showed significant growth of {mom:.2f}% compared to the previous month, indicating strong momentum."
            insights.append(insight)
            print(f"  ✓ Added insight: Significant growth in {code}")
        
        # Rule 2: Notable month-over-month decline (< -3%)
        if mom is not None and mom < -3:
            insight = f"{name} decreased by {abs(mom):.2f}% from the previous month, which requires attention and further analysis."
            insights.append(insight)
            print(f"  ✓ Added insight: Notable decline in {code}")
        
        # Rule 3: Strong year-over-year performance (>8%)
        if yoy is not None and yoy > 8:
            insight = f"{name} demonstrated strong annual performance with {yoy:.2f}% growth compared to the same period last year."
            insights.append(insight)
            print(f"  ✓ Added insight: Strong annual growth in {code}")
        
        # Rule 4: Specific insights for key indicators
        if code == 'UNEMPLOYMENT' and mom is not None and mom < -5:
            insight = f"Unemployment rate fell significantly by {abs(mom):.2f}%, signaling a strengthening labor market."
            insights.append(insight)
            print(f"  ✓ Added insight: Labor market improvement")
        
        if code == 'CPI' and yoy is not None and yoy > 2 and yoy < 4:
            insight = f"Inflation remains moderate at {yoy:.2f}% year-over-year, within the target range."
            insights.append(insight)
            print(f"  ✓ Added insight: Moderate inflation")
    
    # If no specific insights generated, add a general one
    if len(insights) == 0:
        insights.append("Economic indicators showed mixed performance in the reporting period, with some sectors showing growth while others remained stable.")
        print("  ℹ Added default insight")
    
    print(f"✓ Generated {len(insights)} total insights")
    
    return insights


def populate_excel_report(metrics, template_path, output_path):
    """
    Populate the Excel template with calculated metrics
    
    Args:
        metrics (dict): Dictionary of metrics
        template_path (str): Path to Excel template
        output_path (str): Where to save the output file
    """
    try:
        print("\nCreating Excel report...")
        
        # Load the Excel template
        workbook = openpyxl.load_workbook(template_path)
        sheet = workbook['Summary']
        
        # Write the report title and date
        sheet['B2'] = "Monthly Statistics Report"
        sheet['B3'] = f"Period: October 2024"
        
        # Format headers
        sheet['B2'].font = Font(size=16, bold=True)
        sheet['B3'].font = Font(size=12, italic=True)
        
        # Write metrics to the table (starting at row 5)
        row_num = 5
        
        # Define the order of indicators
        indicator_order = ['GDP', 'CPI', 'UNEMPLOYMENT', 'RETAIL_SALES', 
                          'INDUSTRIAL_PROD', 'HOUSING_STARTS', 'EXPORTS', 'IMPORTS']
        
        for code in indicator_order:
            if code in metrics:
                m = metrics[code]
                
                # Column B: Indicator name
                sheet[f'B{row_num}'] = m['name']
                
                # Column C: Current value
                sheet[f'C{row_num}'] = m['current_value']
                
                # Column D: Previous month value
                sheet[f'D{row_num}'] = m['previous_month_value']
                
                # Column E has formula =(C-D)/D which will auto-calculate
                # We can also write it explicitly:
                if m['previous_month_value'] and m['previous_month_value'] != 0:
                    change_pct = ((m['current_value'] - m['previous_month_value']) / 
                                 m['previous_month_value'])
                    sheet[f'E{row_num}'] = change_pct
                    sheet[f'E{row_num}'].number_format = '0.00%'
                
                # Format cells
                sheet[f'C{row_num}'].number_format = '#,##0.00'
                sheet[f'D{row_num}'].number_format = '#,##0.00'
                
                # Color code the change percentage
                if m['mom_growth'] and m['mom_growth'] > 0:
                    sheet[f'E{row_num}'].font = Font(color="00008000")  # Green
                elif m['mom_growth'] and m['mom_growth'] < 0:
                    sheet[f'E{row_num}'].font = Font(color="00800000")  # Red
                
                row_num += 1
        
        print(f"  ✓ Populated {len(indicator_order)} indicators")
        
        # Add insights (optional enhancement)
        insights = generate_insights(metrics)
        insight_row = 15
        for i, insight in enumerate(insights[:6]):  # Max 6 insights
            sheet[f'B{insight_row + i}'] = f"• {insight}"
            sheet[f'B{insight_row + i}'].alignment = Alignment(wrap_text=True)
        
        # Update chart if it exists (optional)
        # This would require the Historical Data sheet to be populated
        
        # Save the workbook
        workbook.save(output_path)
        print(f"✓ Excel report saved: {output_path}")
        
    except FileNotFoundError:
        print(f"✗ Error: Template file not found: {template_path}")
        print("  Creating a basic report without template...")
        create_basic_excel_report(metrics, output_path)
    except Exception as e:
        print(f"✗ Error creating Excel report: {str(e)}")
        import traceback
        traceback.print_exc()


def create_basic_excel_report(metrics, output_path):
    """
    Create a basic Excel report without a template (fallback)
    """
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Summary"
    
    # Headers
    sheet['A1'] = "Monthly Statistics Report"
    sheet['A2'] = f"Period: October 2024"
    sheet['A1'].font = Font(size=16, bold=True)
    
    # Table headers
    headers = ['Indicator', 'Current Value', 'Previous Month', 'Change %']
    for col, header in enumerate(headers, start=1):
        cell = sheet.cell(row=4, column=col)
        cell.value = header
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="DDDDDD", end_color="DDDD DD", fill_type="solid")
    
    # Data
    row_num = 5
    for code, m in metrics.items():
        sheet.cell(row=row_num, column=1).value = m['name']
        sheet.cell(row=row_num, column=2).value = m['current_value']
        sheet.cell(row=row_num, column=3).value = m['previous_month_value']
        if m['mom_growth']:
            sheet.cell(row=row_num, column=4).value = m['mom_growth'] / 100
            sheet.cell(row=row_num, column=4).number_format = '0.00%'
        row_num += 1
    
    workbook.save(output_path)
    print(f"✓ Basic Excel report created: {output_path}")


def create_press_release(metrics, insights, output_path):
    """
    Create a press release document
    
    Args:
        metrics (dict): Dictionary of metrics
        insights (list): List of insight strings
        output_path (str): Where to save the document
    """
    try:
        print("\nCreating press release...")
        
        # Create a new Word document
        doc = Document()
        
        # Add header
        header = doc.add_heading('DEPARTMENT OF STATISTICS', 0)
        header.alignment = 1  # Center alignment
        
        subheader = doc.add_heading('PRESS RELEASE', level=2)
        subheader.alignment = 1
        
        # Add date
        date_para = doc.add_paragraph()
        date_para.add_run(f'Release Date: {datetime.now().strftime("%B %d, %Y")}')
        date_para.alignment = 1
        
        # Add spacing
        doc.add_paragraph()
        
        # Add title
        doc.add_heading('Monthly Economic Indicators - October 2024', level=1)
        
        # Generate summary paragraph using key metrics
        gdp_data = metrics.get('GDP', {})
        cpi_data = metrics.get('CPI', {})
        unemp_data = metrics.get('UNEMPLOYMENT', {})
        
        summary_parts = []
        
        if gdp_data.get('mom_growth'):
            gdp_change = "increased" if gdp_data['mom_growth'] > 0 else "decreased"
            summary_parts.append(
                f"GDP {gdp_change} by {abs(gdp_data['mom_growth']):.2f}% to "
                f"{gdp_data['current_value']:,.0f} million"
            )
        
        if cpi_data.get('yoy_growth'):
            summary_parts.append(
                f"inflation stood at {cpi_data['yoy_growth']:.2f}% year-over-year"
            )
        
        if unemp_data.get('current_value'):
            summary_parts.append(
                f"the unemployment rate was {unemp_data['current_value']:.1f}%"
            )
        
        summary = (
            f"In October 2024, {', '.join(summary_parts[:-1])}, and {summary_parts[-1]}. "
            f"These indicators reflect the current state of the economy and provide "
            f"insights into economic trends and performance."
        )
        
        summary_para = doc.add_paragraph(summary)
        summary_para.alignment = 3  # Justify
        
        # Add Key Highlights section
        doc.add_heading('Key Highlights', level=2)
        
        # Add insights as bullet points
        for insight in insights:
            doc.add_paragraph(insight, style='List Bullet')
        
        # Add additional context
        doc.add_paragraph()
        context = doc.add_paragraph(
            "The Department of Statistics continues to monitor economic indicators "
            "and will provide updated data in the next monthly release. "
            "Detailed statistical tables and historical data are available on our website."
        )
        context.alignment = 3
        
        # Add footer
        doc.add_paragraph()
        doc.add_paragraph('─' * 60)
        
        contact_heading = doc.add_paragraph()
        contact_heading.add_run('For More Information:').bold = True
        
        doc.add_paragraph('Department of Statistics')
        doc.add_paragraph('Email: statistics@gov.example')
        doc.add_paragraph('Website: www.statistics.gov.example')
        
        # Save the document
        doc.save(output_path)
        print(f"✓ Press release saved: {output_path}")
        print(f"  Word count: ~{len(summary.split()) + sum(len(i.split()) for i in insights)} words")
        
    except Exception as e:
        print(f"✗ Error creating press release: {str(e)}")
        import traceback
        traceback.print_exc()


def create_simple_dashboard(data, metrics):
    """
    Create a simple interactive dashboard using Plotly
    
    This is a BONUS implementation showing advanced visualization
    """
    try:
        import plotly.graph_objects as go
        from plotly.subplots import make_subplots
        
        print("\nCreating interactive dashboard...")
        
        # Load historical data for trends
        try:
            historical = pd.read_csv('data/historical_trends.csv')
        except:
            print("  ℹ Historical data not available, creating simplified dashboard")
            historical = None
        
        # Create subplots: 2 rows, 2 columns
        fig = make_subplots(
            rows=2, cols=2,
            subplot_titles=('Historical Trends (GDP & CPI)', 
                          'Month-over-Month Change',
                          'Year-over-Year Comparison',
                          'Key Metrics Summary'),
            specs=[[{'type': 'scatter'}, {'type': 'bar'}],
                   [{'type': 'bar'}, {'type': 'indicator'}]]
        )
        
        # Chart 1: Historical Trends (if data available)
        if historical is not None:
            for indicator in ['GDP', 'CPI']:
                hist_data = historical[historical['indicator_code'] == indicator]
                fig.add_trace(
                    go.Scatter(
                        x=hist_data['period'],
                        y=hist_data['value'],
                        name=indicator,
                        mode='lines+markers'
                    ),
                    row=1, col=1
                )
        
        # Chart 2: Month-over-Month Change
        mom_indicators = []
        mom_values = []
        mom_colors = []
        
        for code, m in metrics.items():
            if m['mom_growth'] is not None:
                mom_indicators.append(code)
                mom_values.append(m['mom_growth'])
                mom_colors.append('green' if m['mom_growth'] > 0 else 'red')
        
        fig.add_trace(
            go.Bar(
                x=mom_indicators,
                y=mom_values,
                name='MoM Change %',
                marker_color=mom_colors
            ),
            row=1, col=2
        )
        
        # Chart 3: Year-over-Year Comparison
        yoy_indicators = []
        yoy_values = []
        
        for code, m in metrics.items():
            if m['yoy_growth'] is not None:
                yoy_indicators.append(code)
                yoy_values.append(m['yoy_growth'])
        
        fig.add_trace(
            go.Bar(
                x=yoy_indicators,
                y=yoy_values,
                name='YoY Change %',
                marker_color='steelblue'
            ),
            row=2, col=1
        )
        
        # Chart 4: Key Metric Indicator (GDP)
        gdp_value = metrics.get('GDP', {}).get('current_value', 0)
        gdp_delta = metrics.get('GDP', {}).get('mom_growth', 0)
        
        fig.add_trace(
            go.Indicator(
                mode="number+delta",
                value=gdp_value,
                title={'text': "GDP (Millions)"},
                delta={'reference': metrics.get('GDP', {}).get('previous_month_value', 0),
                       'relative': False},
                domain={'x': [0, 1], 'y': [0, 1]}
            ),
            row=2, col=2
        )
        
        # Update layout
        fig.update_layout(
            title_text="Monthly Economic Dashboard - October 2024",
            showlegend=True,
            height=800,
            template='plotly_white'
        )
        
        # Save to HTML
        output_file = 'output/dashboard.html'
        fig.write_html(output_file)
        print(f"✓ Dashboard saved: {output_file}")
        print(f"  Open {output_file} in your browser to view the interactive dashboard")
        
    except ImportError:
        print("ℹ Plotly not installed - skipping dashboard creation")
        print("  To install: pip install plotly")
    except Exception as e:
        print(f"✗ Error creating dashboard: {str(e)}")
        import traceback
        traceback.print_exc()


def main():
    """Main workflow orchestration"""
    
    print("="*80)
    print("Exercise 2 Solution: Automated Report Generation")
    print("="*80)
    print()
    
    # Ensure output directory exists
    os.makedirs('output', exist_ok=True)
    print("✓ Output directory ready")
    print()
    
    # Step 1: Load data
    print("STEP 1: Loading data")
    print("-" * 80)
    data = load_monthly_data()
    if data is None:
        print("✗ Cannot proceed without data")
        return
    
    # Step 2: Calculate metrics
    print("\n" + "="*80)
    print("STEP 2: Calculating metrics")
    print("-" * 80)
    metrics = calculate_metrics(data)
    print(f"\n✓ Calculated metrics for {len(metrics)} indicators")
    
    # Step 3: Generate insights
    print("\n" + "="*80)
    print("STEP 3: Generating insights")
    print("-" * 80)
    insights = generate_insights(metrics)
    print(f"\n✓ Generated {len(insights)} insights")
    
    # Step 4: Create Excel report
    print("\n" + "="*80)
    print("STEP 4: Creating Excel report")
    print("-" * 80)
    populate_excel_report(
        metrics,
        'templates/monthly_report_template.xlsx',
        'output/monthly_report_Oct2024.xlsx'
    )
    
    # Step 5: Create press release
    print("\n" + "="*80)
    print("STEP 5: Creating press release")
    print("-" * 80)
    create_press_release(
        metrics,
        insights,
        'output/press_release_Oct2024.docx'
    )
    
    # Step 6 (Bonus): Create dashboard
    print("\n" + "="*80)
    print("STEP 6 (BONUS): Creating interactive dashboard")
    print("-" * 80)
    create_simple_dashboard(data, metrics)
    
    # Summary
    print("\n" + "="*80)
    print("WORKFLOW COMPLETE!")
    print("="*80)
    print("\nGenerated files:")
    print("  📊 output/monthly_report_Oct2024.xlsx")
    print("  📄 output/press_release_Oct2024.docx")
    print("  📈 output/dashboard.html (if Plotly installed)")
    
    print("\nKey Achievements:")
    print(f"  ✓ Processed {len(data)} indicators")
    print(f"  ✓ Calculated {len(metrics) * 2} growth metrics (MoM & YoY)")
    print(f"  ✓ Generated {len(insights)} data-driven insights")
    print(f"  ✓ Created 2-3 publication-ready outputs")
    
    print("\nNext Steps:")
    print("  1. Open the Excel file and verify calculations")
    print("  2. Review the press release for content and tone")
    print("  3. View the dashboard in your browser (if created)")
    print("  4. Consider: How could this be scheduled monthly?")
    print("  5. Think about: What additional metrics would be valuable?")
    
    print("\n" + "="*80)


if __name__ == "__main__":
    main()