"""
Demo 2: Automated Report Generation
Instructor demonstration showing end-to-end automation workflow

This demo is designed for live presentation and includes:
- Step-by-step walkthrough of automation workflow
- Visual progress indicators
- Sample outputs at each stage
- Clear explanations of what's happening
"""

import pandas as pd
import openpyxl
from openpyxl.styles import Font, PatternFill
from docx import Document
from docx.shared import Pt
from datetime import datetime
import os
import time


class Demo:
    """Helper class for demo presentation"""
    
    @staticmethod
    def pause(message="Press Enter to continue...", seconds=0):
        """Pause for audience comprehension"""
        if seconds > 0:
            time.sleep(seconds)
        else:
            input(f"\n{message}")
    
    @staticmethod
    def section(title):
        """Print a section header"""
        print("\n" + "="*80)
        print(f"  {title}")
        print("="*80 + "\n")
    
    @staticmethod
    def subsection(title):
        """Print a subsection header"""
        print("\n" + "─"*80)
        print(f"  {title}")
        print("─"*80 + "\n")


def demo_manual_workflow():
    """
    DEMO PART 1: Show the pain of manual workflows
    Illustrate what we're trying to automate
    """
    Demo.section("PART 1: The Current Manual Workflow (The Problem)")
    
    print("Let me show you what a typical monthly reporting process looks like...\n")
    
    steps = [
        ("Day 1, 9:00 AM", "Export data from 3 different systems"),
        ("Day 1, 10:30 AM", "Open Excel template, start copy-pasting"),
        ("Day 1, 11:45 AM", "Realize you pasted into wrong column, start over"),
        ("Day 1, 2:00 PM", "Manually update all formulas and charts"),
        ("Day 1, 4:30 PM", "Write narrative text for press release"),
        ("Day 2, 9:00 AM", "Manager finds error, need to redo everything"),
        ("Day 2, 11:00 AM", "Fix errors, update charts again"),
        ("Day 2, 3:00 PM", "Format press release in Word"),
        ("Day 2, 4:45 PM", "Final review, find more typos"),
        ("Day 3, 10:00 AM", "Finally send to stakeholders"),
    ]
    
    print("Timeline of Manual Report Generation:\n")
    
    for timestamp, task in steps:
        print(f"  {timestamp:20s} → {task}")
        time.sleep(0.5)
    
    print("\n" + "─"*80)
    print(f"  TOTAL TIME: ~2.5 days (20 hours)")
    print(f"  ERRORS INTRODUCED: Multiple (copy-paste, formula mistakes)")
    print(f"  STRESS LEVEL: High (especially near deadline)")
    print("─"*80)
    
    Demo.pause()
    
    print("\n🎯 THE GOAL:")
    print("   Reduce 2.5 days → 2 minutes")
    print("   Eliminate errors from manual processes")
    print("   Free staff to do actual analysis, not data entry")
    
    Demo.pause()


def demo_data_loading():
    """
    DEMO PART 2: Load and explore data
    Show what the input data looks like
    """
    Demo.section("PART 2: Loading Data (Step 1)")
    
    print("First, let's load our monthly statistics data...\n")
    
    Demo.pause()
    
    print("Loading data/sample_monthly_data.csv...")
    data = pd.read_csv('data/sample_monthly_data.csv')
    
    print(f"\n✓ Loaded {len(data)} indicators\n")
    
    print("Here's what the data looks like:\n")
    print(data[['indicator_code', 'indicator_name', 'value', 'previous_month_value']].to_string(index=False))
    
    Demo.pause()
    
    print("\nNotice we have:")
    print(f"  • Current values (period: {data['period'].iloc[0]})")
    print(f"  • Previous month values (for month-over-month comparison)")
    print(f"  • Previous year values (for year-over-year comparison)")
    
    print("\nThis is everything we need to calculate growth rates and trends!")
    
    Demo.pause()
    
    return data


def demo_metric_calculation(data):
    """
    DEMO PART 3: Calculate metrics
    Show automated calculations
    """
    Demo.section("PART 3: Calculating Metrics (Step 2)")
    
    print("Now let's calculate growth rates for each indicator...\n")
    
    Demo.pause()
    
    print("Calculating month-over-month and year-over-year growth:\n")
    
    metrics = {}
    
    for idx, row in data.iterrows():
        code = row['indicator_code']
        name = row['indicator_name']
        current = row['value']
        prev_month = row['previous_month_value']
        prev_year = row['previous_year_value']
        
        # Calculate MoM growth
        if prev_month and prev_month != 0:
            mom_growth = ((current - prev_month) / prev_month) * 100
        else:
            mom_growth = None
        
        # Calculate YoY growth
        if prev_year and prev_year != 0:
            yoy_growth = ((current - prev_year) / prev_year) * 100
        else:
            yoy_growth = None
        
        # Determine trend
        if mom_growth is not None:
            if mom_growth > 1.0:
                trend = "↑ increasing"
            elif mom_growth < -1.0:
                trend = "↓ decreasing"
            else:
                trend = "→ stable"
        else:
            trend = "? unknown"
        
        metrics[code] = {
            'name': name,
            'current': current,
            'mom_growth': mom_growth,
            'yoy_growth': yoy_growth,
            'trend': trend
        }
        
        # Display calculation
        print(f"  {code:20s} MoM: {mom_growth:>6.2f}%  YoY: {yoy_growth:>6.2f}%  {trend}")
        time.sleep(0.3)
    
    print(f"\n✓ Calculated metrics for {len(metrics)} indicators")
    
    Demo.pause()
    
    print("\n🔍 WHAT JUST HAPPENED:")
    print("   In seconds, we calculated what would take")
    print("   30+ minutes manually (and with no errors!)")
    
    Demo.pause()
    
    return metrics


def demo_insight_generation(metrics):
    """
    DEMO PART 4: Generate insights
    Show rule-based insight generation
    """
    Demo.section("PART 4: Generating Insights (Step 3)")
    
    print("Now let's automatically generate insights from the data...\n")
    print("Using rule-based logic:\n")
    
    rules = [
        "• If MoM growth > 5% → 'Significant growth'",
        "• If MoM growth < -3% → 'Notable decline'",
        "• If YoY growth > 8% → 'Strong annual performance'",
        "• Specific rules for key indicators (e.g., unemployment)"
    ]
    
    for rule in rules:
        print(f"  {rule}")
    
    Demo.pause()
    
    print("\nApplying rules to our data:\n")
    
    insights = []
    
    for code, m in metrics.items():
        mom = m['mom_growth']
        yoy = m['yoy_growth']
        name = m['name']
        
        # Rule 1: Significant MoM growth
        if mom is not None and mom > 5:
            insight = f"{name} showed significant growth of {mom:.2f}% compared to the previous month."
            insights.append(insight)
            print(f"✓ Rule matched: Significant growth in {code}")
            print(f"  → Generated: '{insight}'\n")
            time.sleep(1)
        
        # Rule 2: Notable decline
        if mom is not None and mom < -3:
            insight = f"{name} decreased by {abs(mom):.2f}% from the previous month."
            insights.append(insight)
            print(f"✓ Rule matched: Notable decline in {code}")
            print(f"  → Generated: '{insight}'\n")
            time.sleep(1)
        
        # Rule 3: Strong YoY
        if yoy is not None and yoy > 8:
            insight = f"{name} demonstrated strong annual performance with {yoy:.2f}% year-over-year growth."
            insights.append(insight)
            print(f"✓ Rule matched: Strong YoY in {code}")
            print(f"  → Generated: '{insight}'\n")
            time.sleep(1)
    
    if len(insights) == 0:
        insights.append("Economic indicators showed mixed performance in the reporting period.")
        print("No strong signals detected, using default insight\n")
    
    print(f"✓ Generated {len(insights)} insights")
    
    Demo.pause()
    
    print("\n🔍 WHY THIS MATTERS:")
    print("   These insights are:")
    print("   • Consistent in tone and style")
    print("   • Data-driven (no human bias)")
    print("   • Generated in seconds, not hours")
    print("   • Can be refined by updating rules")
    
    Demo.pause()
    
    return insights


def demo_excel_creation(metrics):
    """
    DEMO PART 5: Create Excel report
    Show automated Excel generation
    """
    Demo.section("PART 5: Creating Excel Report (Step 4)")
    
    print("Now let's create the Excel report...\n")
    
    Demo.pause()
    
    print("Creating new workbook with data and formatting...")
    
    # Create workbook
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Summary"
    
    # Add title
    ws['A1'] = "Monthly Statistics Report"
    ws['A1'].font = Font(size=16, bold=True)
    ws['A2'] = f"Period: October 2024"
    ws['A2'].font = Font(size=12, italic=True)
    
    print("  ✓ Added title and date")
    time.sleep(0.5)
    
    # Add headers
    headers = ['Indicator', 'Current Value', 'Previous Month', 'Change %']
    for col, header in enumerate(headers, start=1):
        cell = ws.cell(row=4, column=col)
        cell.value = header
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        cell.font = Font(bold=True, color="FFFFFF")
    
    print("  ✓ Added column headers with formatting")
    time.sleep(0.5)
    
    # Add data
    print("\n  Adding indicator data:")
    row_num = 5
    for code, m in metrics.items():
        ws.cell(row=row_num, column=1).value = m['name']
        ws.cell(row=row_num, column=2).value = m['current']
        
        if m['mom_growth'] is not None:
            ws.cell(row=row_num, column=4).value = m['mom_growth'] / 100
            ws.cell(row=row_num, column=4).number_format = '0.00%'
            
            # Color code
            if m['mom_growth'] > 0:
                ws.cell(row=row_num, column=4).font = Font(color="008000")
            else:
                ws.cell(row=row_num, column=4).font = Font(color="800000")
        
        print(f"    → {m['name']:40s} {m['mom_growth']:>6.2f}%")
        row_num += 1
        time.sleep(0.2)
    
    # Save
    os.makedirs('output', exist_ok=True)
    output_path = 'output/demo_monthly_report.xlsx'
    wb.save(output_path)
    
    print(f"\n✓ Excel report saved: {output_path}")
    
    Demo.pause()
    
    print("\n🔍 WHAT WE CREATED:")
    print("   • Professional formatted report")
    print("   • All calculations done automatically")
    print("   • Color-coded changes (green/red)")
    print("   • Ready to share with stakeholders")
    
    Demo.pause()


def demo_word_creation(metrics, insights):
    """
    DEMO PART 6: Create press release
    Show automated Word document generation
    """
    Demo.section("PART 6: Creating Press Release (Step 5)")
    
    print("Finally, let's create the press release...\n")
    
    Demo.pause()
    
    print("Creating Word document...")
    
    # Create document
    doc = Document()
    
    # Header
    header = doc.add_heading('DEPARTMENT OF STATISTICS', 0)
    header.alignment = 1
    
    subheader = doc.add_heading('PRESS RELEASE', level=2)
    subheader.alignment = 1
    
    print("  ✓ Added header")
    time.sleep(0.5)
    
    # Date
    date_para = doc.add_paragraph()
    date_para.add_run(f'Release Date: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = 1
    
    print("  ✓ Added date")
    time.sleep(0.5)
    
    # Title
    doc.add_heading('Monthly Economic Indicators - October 2024', level=1)
    
    print("  ✓ Added title")
    time.sleep(0.5)
    
    # Summary paragraph
    gdp = metrics.get('GDP', {})
    cpi = metrics.get('CPI', {})
    unemp = metrics.get('UNEMPLOYMENT', {})
    
    summary = (
        f"In October 2024, GDP {'increased' if gdp.get('mom_growth', 0) > 0 else 'decreased'} "
        f"by {abs(gdp.get('mom_growth', 0)):.2f}% to {gdp.get('current', 0):,.0f} million, "
        f"inflation stood at {cpi.get('yoy_growth', 0):.2f}% year-over-year, "
        f"and the unemployment rate was {unemp.get('current', 0):.1f}%. "
        f"These indicators reflect the current state of the economy."
    )
    
    doc.add_paragraph(summary)
    
    print("  ✓ Added summary paragraph")
    print(f"\n  Summary preview:")
    print(f"  '{summary[:100]}...'\n")
    time.sleep(1)
    
    # Key highlights
    doc.add_heading('Key Highlights', level=2)
    
    print("  Adding key highlights:")
    for insight in insights:
        doc.add_paragraph(insight, style='List Bullet')
        print(f"    • {insight[:70]}...")
        time.sleep(0.5)
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph('─' * 60)
    doc.add_paragraph('Department of Statistics')
    doc.add_paragraph('Email: statistics@gov.example')
    
    print("\n  ✓ Added contact information")
    time.sleep(0.5)
    
    # Save
    output_path = 'output/demo_press_release.docx'
    doc.save(output_path)
    
    print(f"\n✓ Press release saved: {output_path}")
    
    Demo.pause()
    
    print("\n🔍 WHAT WE CREATED:")
    print("   • Professional press release")
    print("   • Data-driven narrative")
    print("   • Consistent formatting")
    print("   • Ready for distribution")
    
    Demo.pause()


def demo_complete_workflow():
    """
    DEMO PART 7: Show the complete workflow
    Tie everything together
    """
    Demo.section("PART 7: Complete Workflow in Action")
    
    print("Let's run the ENTIRE workflow from start to finish...\n")
    print("Watch how fast this happens!\n")
    
    Demo.pause("Press Enter to start the automated workflow...")
    
    start_time = time.time()
    
    # Step 1
    print("\n[Step 1/5] Loading data...")
    data = pd.read_csv('data/sample_monthly_data.csv')
    print(f"  ✓ Loaded {len(data)} indicators")
    time.sleep(0.5)
    
    # Step 2
    print("\n[Step 2/5] Calculating metrics...")
    metrics = {}
    for idx, row in data.iterrows():
        code = row['indicator_code']
        mom = ((row['value'] - row['previous_month_value']) / row['previous_month_value']) * 100
        yoy = ((row['value'] - row['previous_year_value']) / row['previous_year_value']) * 100
        metrics[code] = {
            'name': row['indicator_name'],
            'current': row['value'],
            'mom_growth': mom,
            'yoy_growth': yoy,
            'trend': "increasing" if mom > 1 else "decreasing" if mom < -1 else "stable"
        }
    print(f"  ✓ Calculated {len(metrics)} indicators")
    time.sleep(0.5)
    
    # Step 3
    print("\n[Step 3/5] Generating insights...")
    insights = []
    for code, m in metrics.items():
        if m['mom_growth'] > 5:
            insights.append(f"{m['name']} showed significant growth of {m['mom_growth']:.2f}%.")
    if not insights:
        insights.append("Economic indicators showed mixed performance.")
    print(f"  ✓ Generated {len(insights)} insights")
    time.sleep(0.5)
    
    # Step 4
    print("\n[Step 4/5] Creating Excel report...")
    wb = openpyxl.Workbook()
    ws = wb.active
    ws['A1'] = "Monthly Statistics Report"
    row = 5
    for code, m in metrics.items():
        ws.cell(row=row, column=1).value = m['name']
        ws.cell(row=row, column=2).value = m['current']
        row += 1
    wb.save('output/demo_report.xlsx')
    print(f"  ✓ Saved Excel report")
    time.sleep(0.5)
    
    # Step 5
    print("\n[Step 5/5] Creating press release...")
    doc = Document()
    doc.add_heading('Monthly Statistics - October 2024', 0)
    doc.add_paragraph("Economic summary...")
    for insight in insights:
        doc.add_paragraph(insight, style='List Bullet')
    doc.save('output/demo_press_release.docx')
    print(f"  ✓ Saved press release")
    time.sleep(0.5)
    
    elapsed = time.time() - start_time
    
    print("\n" + "="*80)
    print(f"  WORKFLOW COMPLETE!")
    print("="*80)
    print(f"\n  Total time: {elapsed:.1f} seconds")
    print(f"  Files created:")
    print(f"    • output/demo_report.xlsx")
    print(f"    • output/demo_press_release.docx")
    
    print("\n  Compare to manual process:")
    print(f"    Manual: ~2.5 days (20 hours)")
    print(f"    Automated: {elapsed:.1f} seconds")
    print(f"    Time saved: 99.99%")
    
    Demo.pause()


def demo_real_world_deployment():
    """
    DEMO PART 8: Discuss production deployment
    Share practical deployment advice
    """
    Demo.section("PART 8: Deploying to Production")
    
    print("How to deploy this in your actual workflow:\n")
    
    steps = [
        ("Schedule it", 
         "Use cron/Task Scheduler to run on 1st of month"),
        
        ("Add error handling",
         "What happens if data is missing or looks wrong?"),
        
        ("Set up notifications",
         "Email the report automatically when done"),
        
        ("Create a review step",
         "Human review before final distribution"),
        
        ("Version control outputs",
         "Archive previous months for audit trail"),
        
        ("Monitor and alert",
         "Get notified if the workflow fails"),
        
        ("Document the process",
         "So others can maintain it when you're on vacation"),
        
        ("Start simple",
         "Automate one report first, then expand"),
    ]
    
    print("Deployment Checklist:\n")
    
    for i, (step, detail) in enumerate(steps, 1):
        print(f"{i}. {step}")
        print(f"   → {detail}\n")
        time.sleep(0.8)
    
    Demo.pause()
    
    print("\n🎯 REMEMBER:")
    print("   Automation is not 'set and forget'")
    print("   It's 'set, monitor, improve'")
    
    Demo.pause()


def main():
    """Run the complete demonstration"""
    
    print("\n" + "="*80)
    print("  DEMO: Automated Report Generation")
    print("  From Manual Drudgery to Automated Excellence")
    print("  Instructor: [Your Name]")
    print("="*80)
    
    Demo.pause("Press Enter to start the demonstration...")
    
    # Run demonstration
    demo_manual_workflow()
    
    data = demo_data_loading()
    
    metrics = demo_metric_calculation(data)
    
    insights = demo_insight_generation(metrics)
    
    demo_excel_creation(metrics)
    
    demo_word_creation(metrics, insights)
    
    demo_complete_workflow()
    
    demo_real_world_deployment()
    
    # Final summary
    Demo.section("DEMONSTRATION COMPLETE")
    
    print("What we covered:")
    print("  1. The pain of manual workflows (20 hours)")
    print("  2. Loading and exploring data")
    print("  3. Automated metric calculations")
    print("  4. Rule-based insight generation")
    print("  5. Excel report creation")
    print("  6. Press release generation")
    print("  7. Complete end-to-end workflow (<30 seconds)")
    print("  8. Production deployment tips")
    
    print("\nKey Takeaways:")
    print("  ✓ Automation saves massive amounts of time")
    print("  ✓ Reduces errors from manual processes")
    print("  ✓ Provides consistent, high-quality outputs")
    print("  ✓ Frees staff for higher-value work")
    
    print("\nNow it's your turn!")
    print("  → Open exercises/exercise2_report_generation.py")
    print("  → Implement the TODO sections")
    print("  → Generate your own automated reports")
    
    print("\n" + "="*80)
    print("Questions?")
    print("="*80 + "\n")


if __name__ == "__main__":
    # Ensure output directory exists
    os.makedirs('output', exist_ok=True)
    main()