"""
Exercise 2: Automated Report Generation
Create an automated workflow from data to publication

Your Task:
1. Load monthly statistics data
2. Calculate key metrics (growth rates, trends)
3. Generate insights using rule-based logic
4. Populate Excel template with data and charts
5. Create a press release document
"""

import pandas as pd
import openpyxl
from openpyxl.chart import LineChart, Reference
from docx import Document
from datetime import datetime
import os


def load_monthly_data():
    """
    Load the monthly statistics data from CSV
    
    Returns:
        DataFrame: The loaded data
    """
    # This function is complete - no TODO
    try:
        data = pd.read_csv('data/sample_monthly_data.csv')
        print(f"✓ Loaded {len(data)} indicators from CSV")
        return data
    except FileNotFoundError:
        print("✗ Error: data/sample_monthly_data.csv not found")
        print("Make sure you're running from the repository root directory")
        return None


def calculate_metrics(data):
    """
    Calculate key metrics from the data
    
    Args:
        data (DataFrame): The monthly statistics data
        
    Returns:
        dict: Dictionary of calculated metrics
    """
    metrics = {}
    
    # TODO: Calculate metrics for each indicator
    # For each row in the data, calculate:
    # - Month-over-month growth rate
    # - Year-over-year growth rate
    # - Trend direction (increasing/decreasing/stable)
    
    for idx, row in data.iterrows():
        indicator_code = row['indicator_code']
        
        # TODO: Calculate month-over-month growth
        # Formula: ((current - previous_month) / previous_month) * 100
        current_value = row['value']
        previous_month = row['previous_month_value']
        
        mom_growth = None  # TODO: Calculate this
        
        # TODO: Calculate year-over-year growth
        # Formula: ((current - previous_year) / previous_year) * 100
        previous_year = row['previous_year_value']
        
        yoy_growth = None  # TODO: Calculate this
        
        # TODO: Determine trend direction
        # If growth > 1%: "increasing"
        # If growth < -1%: "decreasing"  
        # Otherwise: "stable"
        trend = None  # TODO: Determine this
        
        metrics[indicator_code] = {
            'name': row['indicator_name'],
            'code': indicator_code,
            'current_value': current_value,
            'previous_month_value': previous_month,
            'previous_year_value': previous_year,
            'unit': row['unit'],
            'period': row['period'],
            'mom_growth': mom_growth,
            'yoy_growth': yoy_growth,
            'trend': trend
        }
    
    return metrics


def generate_insights(metrics):
    """
    Generate insights using rule-based conditional logic
    
    Args:
        metrics (dict): Dictionary of calculated metrics
        
    Returns:
        list: List of insight strings
    """
    insights = []
    
    # TODO: Implement rule-based insight generation
    # Rules:
    # - If month-over-month growth > 5%: "Significant growth observed"
    # - If month-over-month growth < -3%: "Notable decline requires attention"
    # - If year-over-year growth > 8%: "Strong annual performance"
    # - For specific indicators, add context
    
    for code, data in metrics.items():
        mom = data['mom_growth']
        yoy = data['yoy_growth']
        name = data['name']
        
        # TODO: Check for significant month-over-month growth
        if mom is not None and mom > 5:
            insight = None  # TODO: Create insight string
            # insights.append(insight)
        
        # TODO: Check for month-over-month decline
        if mom is not None and mom < -3:
            insight = None  # TODO: Create insight string
            # insights.append(insight)
        
        # TODO: Check for strong year-over-year performance
        if yoy is not None and yoy > 8:
            insight = None  # TODO: Create insight string
            # insights.append(insight)
    
    # If no insights generated, add a default one
    if len(insights) == 0:
        insights.append("Economic indicators showed mixed performance in the reporting period.")
    
    return insights


def populate_excel_report(metrics, template_path, output_path):
    """
    Populate the Excel template with calculated metrics
    
    Args:
        metrics (dict): Dictionary of metrics
        template_path (str): Path to Excel template
        output_path (str): Where to save the output file
    """
    try:
        # TODO: Load the Excel template
        # Hint: workbook = openpyxl.load_workbook(template_path)
        workbook = None  # TODO: Load workbook
        
        # TODO: Access the 'Summary' worksheet
        # Hint: sheet = workbook['Summary']
        sheet = None  # TODO: Get worksheet
        
        # TODO: Write the report title and date
        # The template has placeholders in cells B2 and B3
        # sheet['B2'] = "Monthly Statistics Report"
        # sheet['B3'] = f"Period: October 2024"
        
        # TODO: Write metrics to the appropriate cells
        # The template expects data starting at row 5:
        # Column B: Indicator names
        # Column C: Current values
        # Column D: Previous month values
        # Column E: Change % (this has a formula, will auto-calculate)
        
        # Example for first row (row 5):
        # sheet['B5'] = metrics['GDP']['name']
        # sheet['C5'] = metrics['GDP']['current_value']
        # sheet['D5'] = metrics['GDP']['previous_month_value']
        
        # TODO: Loop through metrics and populate cells
        # Start at row 5 and increment for each indicator
        row_num = 5
        for code in ['GDP', 'CPI', 'UNEMPLOYMENT', 'RETAIL_SALES', 
                     'INDUSTRIAL_PROD', 'HOUSING_STARTS', 'EXPORTS', 'IMPORTS']:
            if code in metrics:
                # TODO: Write to cells B, C, D for this row
                pass
                row_num += 1
        
        # TODO: Save the workbook
        # Hint: workbook.save(output_path)
        
        print(f"✓ Excel report created: {output_path}")
        
    except FileNotFoundError:
        print(f"✗ Error: Template file not found: {template_path}")
    except Exception as e:
        print(f"✗ Error creating Excel report: {str(e)}")


def create_press_release(metrics, insights, output_path):
    """
    Create a press release document
    
    Args:
        metrics (dict): Dictionary of metrics
        insights (list): List of insight strings
        output_path (str): Where to save the document
    """
    try:
        # TODO: Create a new Word document
        # Hint: doc = Document()
        doc = None  # TODO: Create document
        
        # TODO: Add title
        # Hint: doc.add_heading('Monthly Statistics Release', 0)
        
        # TODO: Add date
        # Hint: doc.add_paragraph(f'Release Date: {datetime.now().strftime("%B %d, %Y")}')
        
        # TODO: Add summary paragraph using metrics
        # Create a data-driven summary mentioning key indicators
        # Example: "In October 2024, GDP increased by X%, while CPI rose by Y%..."
        
        summary = None  # TODO: Create summary paragraph
        # doc.add_paragraph(summary)
        
        # TODO: Add "Key Highlights" heading
        # Hint: doc.add_heading('Key Highlights', level=2)
        
        # TODO: Add insights as bullet points
        # Hint: for insight in insights:
        #           doc.add_paragraph(insight, style='List Bullet')
        
        # TODO: Add contact information footer
        # doc.add_paragraph('\n---')
        # doc.add_paragraph('Contact: Department of Statistics')
        # doc.add_paragraph('Email: statistics@gov.example')
        
        # TODO: Save the document
        # Hint: doc.save(output_path)
        
        print(f"✓ Press release created: {output_path}")
        
    except Exception as e:
        print(f"✗ Error creating press release: {str(e)}")


def create_simple_dashboard(data, metrics):
    """
    BONUS: Create a simple interactive dashboard using Plotly
    
    This is optional - only attempt if you finish the main exercises
    """
    try:
        import plotly.graph_objects as go
        from plotly.subplots import make_subplots
        
        # TODO: Load historical data for trends
        historical = pd.read_csv('data/historical_trends.csv')
        
        # TODO: Create subplots
        # fig = make_subplots(rows=2, cols=2, subplot_titles=(...))
        
        # TODO: Add trend chart (GDP and CPI over time)
        
        # TODO: Add month-over-month change bar chart
        
        # TODO: Save to HTML
        # fig.write_html('output/dashboard.html')
        
        print("✓ Dashboard created: output/dashboard.html")
        
    except ImportError:
        print("ℹ Plotly not installed - skipping dashboard creation")
    except Exception as e:
        print(f"✗ Error creating dashboard: {str(e)}")


def main():
    """Main workflow"""
    
    print("="*80)
    print("Exercise 2: Automated Report Generation")
    print("="*80)
    print()
    
    # Ensure output directory exists
    os.makedirs('output', exist_ok=True)
    
    # Step 1: Load data
    print("Step 1: Loading data...")
    data = load_monthly_data()
    if data is None:
        print("✗ Cannot proceed without data")
        return
    print()
    
    # Step 2: Calculate metrics
    print("Step 2: Calculating metrics...")
    metrics = calculate_metrics(data)
    print(f"✓ Calculated metrics for {len(metrics)} indicators")
    print()
    
    # Step 3: Generate insights
    print("Step 3: Generating insights...")
    insights = generate_insights(metrics)
    print(f"✓ Generated {len(insights)} insights:")
    for insight in insights:
        print(f"  • {insight}")
    print()
    
    # Step 4: Create Excel report
    print("Step 4: Creating Excel report...")
    populate_excel_report(
        metrics,
        'templates/monthly_report_template.xlsx',
        'output/monthly_report_Oct2024.xlsx'
    )
    print()
    
    # Step 5: Create press release
    print("Step 5: Creating press release...")
    create_press_release(
        metrics,
        insights,
        'output/press_release_Oct2024.docx'
    )
    print()
    
    # Step 6 (Bonus): Create dashboard
    print("Step 6 (Bonus): Creating dashboard...")
    create_simple_dashboard(data, metrics)
    print()
    
    print("="*80)
    print("Exercise Complete!")
    print("="*80)
    print("\nGenerated files:")
    print("  • output/monthly_report_Oct2024.xlsx")
    print("  • output/press_release_Oct2024.docx")
    print("  • output/dashboard.html (if bonus completed)")
    print("\nNext steps:")
    print("1. Open the Excel file and verify the data and calculations")
    print("2. Open the Word document and review the press release")
    print("3. Consider: How could this workflow be scheduled to run automatically?")


if __name__ == "__main__":
    main()