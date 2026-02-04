"""
Exercise 2: Automated Report Generation - COMPLETED EXAMPLE

This is a worked example showing all TODOs completed.
Students can reference this file if they get stuck on any step.

Learning Objectives:
- Process and analyze data programmatically
- Generate Excel reports with charts
- Create formatted Word documents using templates
- Build interactive dashboards
- Automate repetitive reporting tasks

Original file: exercise_02_reports.py
"""

# TODO 1: Import necessary libraries
# COMPLETED: Import all required libraries for data processing and report generation
import pandas as pd
import openpyxl
from openpyxl.chart import LineChart, Reference
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.shared import Pt
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from datetime import datetime
import os


# TODO 2: Load and process monthly data
# COMPLETED: Load CSV data with error handling
def load_monthly_data(file_path='data/sample_monthly_data.csv'):
    """
    Load monthly statistics data from CSV file.

    Args:
        file_path: Path to the CSV file

    Returns:
        DataFrame with monthly statistics
    """
    try:
        df = pd.read_csv(file_path)
        print(f"Loaded {len(df)} indicators from {file_path}")
        print(f"  Indicators: {', '.join(df['indicator_code'].tolist())}")
        return df
    except FileNotFoundError:
        print(f"Error: File not found: {file_path}")
        return None


# TODO 3: Calculate key metrics (Basic)
# COMPLETED: Calculate MoM and YoY growth rates with trend detection
def calculate_metrics(df):
    """
    Calculate growth rates and trends for each indicator.

    Args:
        df: DataFrame with monthly statistics

    Returns:
        Dictionary of metrics for each indicator
    """
    metrics = {}

    print("\nCalculating metrics...")

    for idx, row in df.iterrows():
        indicator_code = row['indicator_code']

        # Extract values
        current_value = row['value']
        previous_month = row['previous_month_value']
        previous_year = row['previous_year_value']

        # Calculate month-over-month growth
        # Formula: ((current - previous) / previous) * 100
        if previous_month and previous_month != 0:
            mom_growth = ((current_value - previous_month) / previous_month) * 100
        else:
            mom_growth = None

        # Calculate year-over-year growth
        # Formula: ((current - previous_year) / previous_year) * 100
        if previous_year and previous_year != 0:
            yoy_growth = ((current_value - previous_year) / previous_year) * 100
        else:
            yoy_growth = None

        # Determine trend direction (threshold: 1% for "stable")
        if mom_growth is not None:
            if mom_growth > 1.0:
                trend = "increasing"
            elif mom_growth < -1.0:
                trend = "decreasing"
            else:
                trend = "stable"
        else:
            trend = "unknown"

        # Store all metrics for this indicator
        metrics[indicator_code] = {
            'name': row['indicator_name'],
            'code': indicator_code,
            'current_value': current_value,
            'previous_month_value': previous_month,
            'previous_year_value': previous_year,
            'unit': row['unit'],
            'period': row['period'],
            'mom_growth': mom_growth,
            'yoy_growth': yoy_growth,
            'trend': trend
        }

        # Print progress
        if mom_growth is not None and yoy_growth is not None:
            print(f"  {indicator_code}: MoM={mom_growth:.2f}%, YoY={yoy_growth:.2f}%, Trend={trend}")

    return metrics


# TODO 3.5: Generate insights from metrics
# COMPLETED: Rule-based insight generation
def generate_insights(df, metrics):
    """
    Generate rule-based insights from data and metrics.

    Args:
        df: Original DataFrame
        metrics: Dictionary of calculated metrics

    Returns:
        List of insight strings
    """
    insights = []

    print("\nGenerating insights...")

    for code, data in metrics.items():
        mom = data['mom_growth']
        yoy = data['yoy_growth']
        name = data['name']

        # Rule 1: Significant month-over-month growth (>5%)
        if mom is not None and mom > 5:
            insight = f"{name} showed significant growth of {mom:.2f}% compared to the previous month."
            insights.append(insight)
            print(f"  + Significant growth in {code}")

        # Rule 2: Notable month-over-month decline (< -3%)
        if mom is not None and mom < -3:
            insight = f"{name} decreased by {abs(mom):.2f}% from the previous month, requiring attention."
            insights.append(insight)
            print(f"  + Notable decline in {code}")

        # Rule 3: Strong year-over-year performance (>8%)
        if yoy is not None and yoy > 8:
            insight = f"{name} demonstrated strong annual performance with {yoy:.2f}% growth year-over-year."
            insights.append(insight)
            print(f"  + Strong annual growth in {code}")

        # Rule 4: Specific insights for unemployment
        if code == 'UNEMPLOYMENT' and mom is not None and mom < -5:
            insight = f"Unemployment rate fell significantly, signaling a strengthening labor market."
            insights.append(insight)
            print(f"  + Labor market improvement")

        # Rule 5: Inflation within target range
        if code == 'CPI' and yoy is not None and 2 < yoy < 4:
            insight = f"Inflation remains moderate at {yoy:.2f}% year-over-year, within target range."
            insights.append(insight)
            print(f"  + Moderate inflation")

    # Default insight if none generated
    if len(insights) == 0:
        insights.append("Economic indicators showed mixed performance in the reporting period.")
        print("  + Added default insight")

    print(f"Generated {len(insights)} insights")
    return insights


# TODO 4: Generate Excel report with charts (Intermediate)
# COMPLETED: Create Excel workbook with data and formatting
def generate_excel_report(df, metrics, output_file='output/monthly_report.xlsx'):
    """
    Create Excel report with data, formatting, and charts.

    Args:
        df: Original DataFrame
        metrics: Dictionary of calculated metrics
        output_file: Path to save the Excel file
    """
    print("\nCreating Excel report...")

    # Create a new workbook
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Summary"

    # Add title
    sheet['A1'] = "Monthly Economic Report"
    sheet['A1'].font = Font(size=16, bold=True)
    sheet['A2'] = f"Period: October 2024"
    sheet['A2'].font = Font(size=12, italic=True)

    # Add table headers
    headers = ['Indicator', 'Current Value', 'Previous Month', 'MoM Change %', 'YoY Change %', 'Trend']
    for col, header in enumerate(headers, start=1):
        cell = sheet.cell(row=4, column=col)
        cell.value = header
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        cell.font = Font(bold=True, color="FFFFFF")

    # Add data rows
    row_num = 5
    for code, m in metrics.items():
        sheet.cell(row=row_num, column=1).value = m['name']
        sheet.cell(row=row_num, column=2).value = m['current_value']
        sheet.cell(row=row_num, column=2).number_format = '#,##0.00'
        sheet.cell(row=row_num, column=3).value = m['previous_month_value']
        sheet.cell(row=row_num, column=3).number_format = '#,##0.00'

        # MoM change with color coding
        if m['mom_growth'] is not None:
            cell = sheet.cell(row=row_num, column=4)
            cell.value = m['mom_growth'] / 100
            cell.number_format = '0.00%'
            if m['mom_growth'] > 0:
                cell.font = Font(color="008000")  # Green
            elif m['mom_growth'] < 0:
                cell.font = Font(color="FF0000")  # Red

        # YoY change
        if m['yoy_growth'] is not None:
            cell = sheet.cell(row=row_num, column=5)
            cell.value = m['yoy_growth'] / 100
            cell.number_format = '0.00%'

        # Trend
        sheet.cell(row=row_num, column=6).value = m['trend']

        row_num += 1

    # Adjust column widths
    sheet.column_dimensions['A'].width = 30
    sheet.column_dimensions['B'].width = 15
    sheet.column_dimensions['C'].width = 15
    sheet.column_dimensions['D'].width = 15
    sheet.column_dimensions['E'].width = 15
    sheet.column_dimensions['F'].width = 12

    # Add a simple bar chart for MoM changes
    # Create chart data sheet
    chart_sheet = workbook.create_sheet("Chart Data")
    chart_sheet['A1'] = "Indicator"
    chart_sheet['B1'] = "MoM Change %"

    chart_row = 2
    for code, m in metrics.items():
        if m['mom_growth'] is not None:
            chart_sheet.cell(row=chart_row, column=1).value = code
            chart_sheet.cell(row=chart_row, column=2).value = m['mom_growth']
            chart_row += 1

    # Save the workbook
    workbook.save(output_file)
    print(f"Excel report saved: {output_file}")


# TODO 5: Create Word document using template (Advanced)
# COMPLETED: Generate formatted Word document
def generate_word_report(metrics, insights, output_file='output/monthly_report.docx'):
    """
    Create formatted Word document with metrics and insights.

    Args:
        metrics: Dictionary of calculated metrics
        insights: List of insight strings
        output_file: Path to save the Word document
    """
    print("\nCreating Word document...")

    # Create a new document
    doc = Document()

    # Add header
    header = doc.add_heading('DEPARTMENT OF STATISTICS', 0)
    header.alignment = 1  # Center

    subheader = doc.add_heading('PRESS RELEASE', level=2)
    subheader.alignment = 1

    # Add date
    date_para = doc.add_paragraph()
    date_para.add_run(f'Release Date: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = 1

    doc.add_paragraph()

    # Add title
    doc.add_heading('Monthly Economic Indicators - October 2024', level=1)

    # Generate summary paragraph
    gdp_data = metrics.get('GDP', {})
    cpi_data = metrics.get('CPI', {})
    unemp_data = metrics.get('UNEMPLOYMENT', {})

    summary_parts = []

    if gdp_data.get('mom_growth'):
        gdp_change = "increased" if gdp_data['mom_growth'] > 0 else "decreased"
        summary_parts.append(
            f"GDP {gdp_change} by {abs(gdp_data['mom_growth']):.2f}% to "
            f"{gdp_data['current_value']:,.0f} million"
        )

    if cpi_data.get('yoy_growth'):
        summary_parts.append(
            f"inflation stood at {cpi_data['yoy_growth']:.2f}% year-over-year"
        )

    if unemp_data.get('current_value'):
        summary_parts.append(
            f"the unemployment rate was {unemp_data['current_value']:.1f}%"
        )

    if summary_parts:
        summary = f"In October 2024, {', '.join(summary_parts[:-1])}, and {summary_parts[-1]}."
    else:
        summary = "Economic indicators for October 2024 are presented below."

    summary_para = doc.add_paragraph(summary)
    summary_para.alignment = 3  # Justify

    # Add Key Highlights section
    doc.add_heading('Key Highlights', level=2)

    for insight in insights:
        doc.add_paragraph(insight, style='List Bullet')

    # Add data table
    doc.add_heading('Summary Statistics', level=2)

    # Create table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Indicator'
    header_cells[1].text = 'Value'
    header_cells[2].text = 'MoM Change'
    header_cells[3].text = 'YoY Change'

    # Data rows
    for code, m in metrics.items():
        row_cells = table.add_row().cells
        row_cells[0].text = m['name']
        row_cells[1].text = f"{m['current_value']:,.2f}"
        row_cells[2].text = f"{m['mom_growth']:.2f}%" if m['mom_growth'] else "N/A"
        row_cells[3].text = f"{m['yoy_growth']:.2f}%" if m['yoy_growth'] else "N/A"

    # Add footer
    doc.add_paragraph()
    doc.add_paragraph('_' * 60)

    contact = doc.add_paragraph()
    contact.add_run('For More Information:').bold = True
    doc.add_paragraph('Department of Statistics')
    doc.add_paragraph('Email: statistics@gov.example')

    # Save the document
    doc.save(output_file)
    print(f"Word document saved: {output_file}")


# TODO 6: Create interactive dashboard (Bonus)
# COMPLETED: Create Plotly dashboard
def generate_dashboard(df, metrics, output_file='output/dashboard.html'):
    """
    Create an interactive dashboard with Plotly.

    Args:
        df: Original DataFrame
        metrics: Dictionary of calculated metrics
        output_file: Path to save the HTML dashboard
    """
    print("\nCreating interactive dashboard...")

    # Create subplots
    fig = make_subplots(
        rows=2, cols=2,
        subplot_titles=(
            'Current Values by Indicator',
            'Month-over-Month Change (%)',
            'Year-over-Year Change (%)',
            'GDP Indicator'
        ),
        specs=[
            [{'type': 'bar'}, {'type': 'bar'}],
            [{'type': 'bar'}, {'type': 'indicator'}]
        ]
    )

    # Prepare data for charts
    indicators = []
    current_values = []
    mom_changes = []
    yoy_changes = []
    mom_colors = []
    yoy_colors = []

    for code, m in metrics.items():
        indicators.append(code)
        current_values.append(m['current_value'])

        if m['mom_growth'] is not None:
            mom_changes.append(m['mom_growth'])
            mom_colors.append('green' if m['mom_growth'] > 0 else 'red')
        else:
            mom_changes.append(0)
            mom_colors.append('gray')

        if m['yoy_growth'] is not None:
            yoy_changes.append(m['yoy_growth'])
            yoy_colors.append('steelblue' if m['yoy_growth'] > 0 else 'orange')
        else:
            yoy_changes.append(0)
            yoy_colors.append('gray')

    # Chart 1: Current Values
    fig.add_trace(
        go.Bar(
            x=indicators,
            y=current_values,
            name='Current Value',
            marker_color='royalblue'
        ),
        row=1, col=1
    )

    # Chart 2: MoM Changes
    fig.add_trace(
        go.Bar(
            x=indicators,
            y=mom_changes,
            name='MoM Change %',
            marker_color=mom_colors
        ),
        row=1, col=2
    )

    # Chart 3: YoY Changes
    fig.add_trace(
        go.Bar(
            x=indicators,
            y=yoy_changes,
            name='YoY Change %',
            marker_color=yoy_colors
        ),
        row=2, col=1
    )

    # Chart 4: GDP Indicator
    gdp_data = metrics.get('GDP', {})
    fig.add_trace(
        go.Indicator(
            mode="number+delta",
            value=gdp_data.get('current_value', 0),
            title={'text': "GDP (Millions)"},
            delta={
                'reference': gdp_data.get('previous_month_value', 0),
                'relative': False
            }
        ),
        row=2, col=2
    )

    # Update layout
    fig.update_layout(
        title_text="Monthly Economic Dashboard - October 2024",
        showlegend=False,
        height=700,
        template='plotly_white'
    )

    # Save to HTML
    fig.write_html(output_file)
    print(f"Dashboard saved: {output_file}")
    print(f"  Open {output_file} in your browser to view")


# TODO 7: Main execution function
# COMPLETED: Orchestrate the entire workflow
def main():
    """
    Main function to run the complete report generation workflow.
    """
    print("=" * 60)
    print("EXERCISE 2: AUTOMATED REPORT GENERATION")
    print("=" * 60)

    # Ensure output directory exists
    os.makedirs('output', exist_ok=True)
    print("\nOutput directory ready")

    # Step 1: Load data
    print("\n" + "-" * 60)
    print("STEP 1: Loading data")
    print("-" * 60)
    df = load_monthly_data()

    if df is None:
        print("Cannot proceed without data")
        return

    # Step 2: Calculate metrics
    print("\n" + "-" * 60)
    print("STEP 2: Calculating metrics")
    print("-" * 60)
    metrics = calculate_metrics(df)
    print(f"\nCalculated metrics for {len(metrics)} indicators")

    # Step 3: Generate insights
    print("\n" + "-" * 60)
    print("STEP 3: Generating insights")
    print("-" * 60)
    insights = generate_insights(df, metrics)

    # Step 4: Generate Excel report
    print("\n" + "-" * 60)
    print("STEP 4: Creating Excel report")
    print("-" * 60)
    generate_excel_report(df, metrics, 'output/monthly_report.xlsx')

    # Step 5: Generate Word report
    print("\n" + "-" * 60)
    print("STEP 5: Creating Word document")
    print("-" * 60)
    generate_word_report(metrics, insights, 'output/monthly_report.docx')

    # Step 6: Generate dashboard (Bonus)
    print("\n" + "-" * 60)
    print("STEP 6 (BONUS): Creating interactive dashboard")
    print("-" * 60)
    generate_dashboard(df, metrics, 'output/dashboard.html')

    # Summary
    print("\n" + "=" * 60)
    print("WORKFLOW COMPLETE!")
    print("=" * 60)
    print("\nGenerated files:")
    print("  - output/monthly_report.xlsx (Excel report)")
    print("  - output/monthly_report.docx (Word document)")
    print("  - output/dashboard.html (Interactive dashboard)")

    print(f"\nSummary:")
    print(f"  - Processed {len(df)} indicators")
    print(f"  - Calculated {len(metrics) * 2} growth metrics (MoM & YoY)")
    print(f"  - Generated {len(insights)} insights")
    print("=" * 60)


if __name__ == "__main__":
    main()
